"""
Base document parser.

Provides shared PDF extraction and file upload handling logic.
Resume and JD parsers inherit from this class and implement
their own section detection strategies.
"""

from __future__ import annotations

import io
import logging
import re
from pathlib import PurePosixPath

from fastapi import UploadFile
from pypdf import PdfReader

from app.config.settings import get_settings
from app.utils.exceptions import (
    EmptyDocumentError,
    FileParsingError,
    FileSizeExceededError,
    UnsupportedFileTypeError,
)

logger = logging.getLogger(__name__)


class BaseDocumentParser:
    """
    Abstract base parser providing PDF extraction and upload handling.

    Subclasses must implement ``parse_sections()`` to detect
    document-specific sections from raw text.
    """

    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """
        Extract raw text from PDF bytes using pypdf.

        Args:
            file_bytes: Raw bytes of the PDF file.

        Returns:
            Concatenated text from all pages.

        Raises:
            FileParsingError: If the PDF is corrupt or unreadable.
            EmptyDocumentError: If no text could be extracted.
        """
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages: list[str] = []

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append(text.strip())
                else:
                    logger.debug("Page %d yielded no text.", page_num + 1)

            full_text = "\n\n".join(pages).strip()

            if not full_text:
                raise EmptyDocumentError("PDF file contains no extractable text.")

            # Automatically un-space spaced character PDFs (e.g. "M O H A M M E D  A D I L  K H A N")
            full_text = self.normalize_spaced_text(full_text)

            logger.debug(
                "Extracted %d chars from %d PDF pages.",
                len(full_text),
                len(reader.pages),
            )
            return full_text

        except (EmptyDocumentError, FileParsingError):
            raise
        except Exception as err:
            logger.error("Failed to parse PDF bytes: %s", err)
            raise FileParsingError(
                f"Failed to parse PDF document: {err}"
            ) from err

    @staticmethod
    def normalize_spaced_text(text: str) -> str:
        """
        Detects and normalizes spaced character text (e.g. 'M O H A M M E D  A D I L  K H A N').
        """
        if len(re.findall(r"\b[A-Za-z0-9]\s[A-Za-z0-9]\s[A-Za-z0-9]\b", text)) > 5:
            cleaned_lines = []
            for line in text.split("\n"):
                line_str = line.strip()
                if not line_str:
                    cleaned_lines.append("")
                    continue
                if len(re.findall(r"\b[A-Za-z0-9]\s[A-Za-z0-9]\b", line_str)) >= 2:
                    tokens = [t.strip() for t in re.split(r"[ \t]{2,}", line_str) if t.strip()]
                    new_tokens = []
                    for tok in tokens:
                        unspaced = re.sub(r"(?<=\b[A-Za-z0-9@.-])\s+(?=[A-Za-z0-9@.-]\b)", "", tok)
                        new_tokens.append(unspaced)
                    line_str = " ".join(new_tokens)
                cleaned_lines.append(line_str)

            t = "\n".join(cleaned_lines)
            # Fix emails with spaced @ or dots
            t = re.sub(r"([a-zA-Z0-9._%+-]+)\s*@\s*([a-zA-Z0-9.-]+)\s*\.\s*([a-zA-Z]{2,})", r"\1@\2.\3", t)
            return t
        return text

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """
        Extract raw text from DOCX bytes using python-docx.

        Extracts text from paragraphs and table cells.

        Args:
            file_bytes: Raw bytes of the DOCX file.

        Returns:
            Concatenated text from all paragraphs and tables.

        Raises:
            FileParsingError: If the DOCX is corrupt or unreadable.
            EmptyDocumentError: If no text could be extracted.
        """
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(file_bytes))

            # Extract paragraph text
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Extract table cell text
            table_texts: list[str] = []
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        table_texts.append(" | ".join(row_cells))

            full_text = "\n".join(paragraphs + table_texts).strip()

            if not full_text:
                raise EmptyDocumentError("DOCX file contains no extractable text.")

            logger.debug(
                "Extracted %d chars from DOCX (%d paragraphs, %d table rows).",
                len(full_text),
                len(paragraphs),
                len(table_texts),
            )
            return full_text

        except (EmptyDocumentError, FileParsingError):
            raise
        except Exception as err:
            logger.error("Failed to parse DOCX bytes: %s", err)
            raise FileParsingError(
                f"Failed to parse DOCX document: {err}"
            ) from err

    async def read_upload(self, file: UploadFile) -> str:
        """
        Read and validate an uploaded file, then extract text.

        Supports PDF and DOCX file formats.

        Args:
            file: The FastAPI UploadFile instance.

        Returns:
            Extracted text content.

        Raises:
            UnsupportedFileTypeError: If the file extension is not allowed.
            FileSizeExceededError: If the file exceeds the size limit.
            FileParsingError: If the file cannot be read or parsed.
        """
        settings = get_settings()

        # Validate file extension
        suffix = ""
        if file.filename:
            suffix = PurePosixPath(file.filename).suffix.lower()
            if suffix not in settings.ALLOWED_EXTENSIONS:
                raise UnsupportedFileTypeError(
                    f"File type '{suffix}' is not supported. "
                    f"Allowed: {', '.join(sorted(settings.ALLOWED_EXTENSIONS))}"
                )

        # Read file bytes
        try:
            file_bytes = await file.read()
        except Exception as exc:
            raise FileParsingError(
                f"Failed to read the uploaded file: {exc}"
            ) from exc

        # Validate file size
        if len(file_bytes) > settings.MAX_FILE_SIZE:
            raise FileSizeExceededError(
                f"File size ({len(file_bytes)} bytes) exceeds the maximum "
                f"allowed limit ({settings.MAX_FILE_SIZE} bytes)."
            )

        if not file_bytes:
            raise EmptyDocumentError("The uploaded file is empty.")

        logger.info(
            "Read uploaded file: %s (%d bytes)",
            file.filename,
            len(file_bytes),
        )

        # Route to correct parser based on extension
        if suffix == ".docx":
            return self.extract_text_from_docx(file_bytes)
        return self.extract_text_from_pdf(file_bytes)

